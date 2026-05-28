"""Refresh resume-source.docx with the website URL, updated skills line,
hyperlinks in the contact header, a trimmed Lynx Equity entry, and a
PERSONAL PROJECTS section (trader + Odds Aggregator).

Idempotent: re-running produces the same docx output. Committed for
reproducibility.

Usage:
    python scripts/refresh-resume.py

In-place rewrites:
    resume-source.docx
"""

from __future__ import annotations

import copy
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "resume-source.docx"
PROJECTS_YML = ROOT / "projects.yml"

# ------------- desired content -------------

# Structured contact line: list of (kind, value-or-(text,url)) tuples.
# 'text' produces a plain run; 'tab' produces a <w:tab/>; 'link' wraps
# a styled run in a <w:hyperlink>.
CONTACT_LINE = [
    ("text", "Toronto, ON"),
    ("tab",  None),
    ("link", ("14jakehoffman@gmail.com", "mailto:14jakehoffman@gmail.com")),
    ("tab",  None),
    ("text", "Tel: 647-823-4504"),
    ("tab",  None),
    ("link", ("LinkedIn",
              "https://www.linkedin.com/in/jake-hoffman-7117692a5/")),
    ("tab",  None),
    ("link", ("jakethehoffer.github.io/website",
              "https://jakethehoffer.github.io/website/")),
]

NEW_SKILLS_LINE = (
    "Proficient in Python, Java, C, C#, C++, SQL, Git, Word, Excel, "
    "JavaScript, Dart, Unity, FastAPI, Playwright, Monday.com, and "
    "Google Apps Script."
)

# Paragraphs to drop from the existing resume body. Matched by exact text prefix.
#
# FROZEN LEDGER (as of v2.22): this is the historical record of one-time
# content removals (e.g. the BMC Pharmacy entry, low-signal Lynx bullets)
# plus page-fit cuts made before the resume_priority + RESUME_MAX_PROJECTS
# cap existed. On the current (already-curated) docx these entries are
# mostly no-ops — the content is already gone.
#
# DO NOT add new page-fit cuts here. Page-fit is now handled declaratively
# by resume_priority + the cap in build_new_section(): adding a project
# competes for the capped slots instead of forcing cuts elsewhere. Add an
# entry here only to permanently exclude a piece of *static* resume
# content (something in SUMMARY/EDUCATION/EXPERIENCE you never want shown).
PARAGRAPHS_TO_DROP = [
    # Lynx Equity bullets (lower-signal items)
    "Designed Monday.com form that prefills",
    "Conducted DNS lookups",
    # BMC Pharmacy entry — drop entirely (pre-Queen's, least
    # aligned with the tech/quant target audience). The trim catches
    # the role line, the org line, and both bullets.
    "Financial Analyst",
    "BMC Pharmacy",
    "Conducted research in SW Ontario",
    "Tracked profitability and margin contribution",
    # PERSONAL PROJECTS — drop redundant subtitles and the most
    # arcane bullet from each project to hit a 1-page resume.
    "AI swing-trading agent",
    "Production cross-book arbitrage daemon",
    "Designed broker abstraction",
    "Built defensive ingestion",
    # QMIND bullet 3 — diluted "helped" phrasing; the first two
    # bullets establish the Q-learning + Deep Q story without it.
    # Dropped in v2.16 to make room for the Writing line.
    "Helped the team move from the simple Q-learning",
    # SUMMARY bullet 4 — generic soft-skill line, lowest signal on
    # the page. Dropped in v2.17 to make room for tax-rebalance.
    "Collaborates with non-technical teammates",
    # trader resume bullet 2 — the website case study tells the
    # kill-switch / risk-gate / paper-trade-gate story in full;
    # the resume reader can click through. Dropped in v2.17.
    "Paper-traded against SPY with kill-switch",
    # OA resume bullet 2 — same reasoning. Website case study covers
    # ingestion + detection in full. Dropped in v2.17.
    "Ingest → normalize → detect cross-book arbs",
    # equity-arbs bullet 1 — old 2-bullet form merged into a single
    # dense bullet in v2.18 to keep the PDF at 1 page. Catches the
    # old form already in the docx. The new merged form starts with
    # "Stat-arb research toolkit for TSX pair trading (spec-first"
    # (parenthesis), which does NOT match this prefix.
    "Stat-arb research toolkit for TSX pair trading. Spec",
    # equity-arbs bullet 2 — old form (now merged into bullet 1).
    "Decision-grade backtest demonstrated",
    # EDUCATION bullet about Walking/Jumping ML — the project is
    # already featured in the website's PROJECTS section, so leaving
    # the resume bullet is a double-count. Dropped in v2.18 to keep
    # the page-fit constraint after adding equity-arbs.
    "Built a machine learning program in Python for accelerometer",
    # EDUCATION bullet about Smart Shoe — same reasoning: the
    # Smart Shoe Navigation project is already featured in the
    # website's PROJECTS section. Dropped in v2.18.
    "Created a mobile app in Dart that pairs via Bluetooth",
    # tax-rebalance old bullet 1 — merged into a single denser bullet
    # in v2.19 to make page-fit room for trades-agency. The new
    # merged form starts with the same prefix, so the trim's exact
    # prefix has to be specific enough to NOT match the new form.
    # Old bullet 1: ends after "never places trades."
    # New merged:   continues into "Spec-driven TDD: 21-task plan..."
    # Use a prefix that catches the old form's exact ending:
    "Canadian TFSA + RRSP portfolio drift monitor — emails weekly digests with cost-aware rebalance verdicts. Read",
    # tax-rebalance old bullet 2 (the spec-driven TDD bullet) —
    # merged into bullet 1.
    "Spec-driven TDD: 21-task plan",
    # trades-agency — dropped from the resume entirely in v2.21
    # (review #3: one negative result on the resume, not two). It was
    # inserted into the docx in v2.19 and walk-anchor only inserts, so
    # the role + bullet are trimmed here. The prefix (no colon) matches
    # both the old and the tightened bullet forms. Stays on the website.
    "trades-agency (Python",
    "Toronto AI-receptionist + lead-gen venture",
    # QMIND bullet 2 — bullet 1 covers the RL story; bullet 2 adds
    # detail at the cost of vertical space. Dropped when adding
    # trades-agency pushed PERSONAL PROJECTS over page boundary.
    "Built a tabular Q-learning version",
    # Mega TTT EXPERIENCE Minimax bullet — the Minimax AI is
    # mentioned on the project's chip list and metrics line on the
    # website, so the resume bullet is redundant. Dropped to make
    # page-fit room for trades-agency.
    "Implemented a Minimax-based algorithm for the computer opponent",
]

# The resume's PERSONAL PROJECTS section shows at most this many
# projects. Selection is by resume_priority (highest first); display
# order is projects.yml order. This is the curation lever that replaces
# the old "include everything, then trim collateral to fit" model —
# adding a project competes for these slots instead of forcing manual
# cuts to EXPERIENCE/EDUCATION. Bump this only if the page genuinely
# has room (verify the PDF stays 1 page).
RESUME_MAX_PROJECTS = 4


def build_new_section() -> list[tuple[str, str]]:
    """Read projects.yml and produce the (kind, text) tuples that the
    walk-anchor logic consumes. The first tuple is always
    ("heading", "PERSONAL PROJECTS").

    Curation: among projects that have a resume: block (candidates),
    select the top RESUME_MAX_PROJECTS by resume_priority (highest
    first), then emit them in projects.yml display order so the resume's
    project order matches the website's. Each selected project produces
    ("role", role) followed by ("bullet", b) for each bullet.

    This makes adding a project safe: a new candidate competes for the
    capped slots by priority rather than forcing page-fit cuts
    elsewhere. The walk-anchor logic in main() handles per-entry
    idempotency."""
    projects = yaml.safe_load(PROJECTS_YML.read_text(encoding="utf-8"))
    candidates = [p for p in projects if p.get("resume")]
    # Select top-N by priority (stable: ties keep projects.yml order).
    ranked = sorted(
        candidates,
        key=lambda p: p.get("resume_priority", 0),
        reverse=True,
    )
    selected_keys = {p["key"] for p in ranked[:RESUME_MAX_PROJECTS]}
    # Emit in projects.yml display order, restricted to the selected set.
    section: list[tuple[str, str]] = [("heading", "PERSONAL PROJECTS")]
    for p in projects:
        if p["key"] not in selected_keys:
            continue
        resume = p["resume"]
        section.append(("role", resume["role"]))
        for bullet in resume.get("bullets", []):
            section.append(("bullet", bullet))
    return section


# Built at import time from projects.yml. Kept as a module-level name
# so the existing main() flow that references NEW_SECTION still works.
NEW_SECTION = build_new_section()

# Writing-section pointer appended at the end of PERSONAL PROJECTS.
# Built as text + hyperlink + text so the URL is clickable in the PDF.
WRITING_LINE = {
    "prefix":    "Writing: ",
    "link_text": "jakethehoffer.github.io/website/#writing",
    "link_url":  "https://jakethehoffer.github.io/website/#writing",
    "suffix":    " — short essays on engineering tradeoffs.",
}


# ------------- helpers -------------

def find_para(paragraphs, predicate):
    for p in paragraphs:
        if predicate(p):
            return p
    return None


def set_paragraph_text(p: Paragraph, new_text: str) -> None:
    """Plain-text replacement preserving the first run's formatting.
    Drops any <w:r> and <w:hyperlink> children."""
    pe = p._element
    first_run = pe.find(qn("w:r"))
    for child in list(pe):
        tag = child.tag.split("}", 1)[-1]
        if tag in ("r", "hyperlink"):
            pe.remove(child)
    if first_run is None:
        first_run = OxmlElement("w:r")
    for t in list(first_run.findall(qn("w:t"))):
        first_run.remove(t)
    for tab in list(first_run.findall(qn("w:tab"))):
        first_run.remove(tab)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    first_run.append(new_t)
    pe.append(first_run)


def clone_paragraph_element(template_p: Paragraph):
    return copy.deepcopy(template_p._element)


def _copy_text_rpr(template_run_element):
    """Return a deep copy of a run's rPr (run properties), or None."""
    if template_run_element is None:
        return None
    rPr = template_run_element.find(qn("w:rPr"))
    if rPr is None:
        return None
    return copy.deepcopy(rPr)


def _make_text_run(text, template_rPr):
    r = OxmlElement("w:r")
    if template_rPr is not None:
        r.append(copy.deepcopy(template_rPr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _make_tab_run(template_rPr):
    r = OxmlElement("w:r")
    if template_rPr is not None:
        r.append(copy.deepcopy(template_rPr))
    r.append(OxmlElement("w:tab"))
    return r


def _make_hyperlink(text, url, template_rPr, header_part):
    """Build a <w:hyperlink> with one styled run inside, and register
    the URL in header_part's relationships."""
    rid = header_part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rid)
    hl.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    # Build rPr: Hyperlink style + selected base properties (font, size, color).
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if template_rPr is not None:
        for child in template_rPr:
            tag = child.tag.split("}", 1)[-1]
            if tag in ("rFonts", "sz", "szCs"):
                rPr.append(copy.deepcopy(child))
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)

    hl.append(r)
    return hl


def build_contact_line(p: Paragraph, header_part) -> None:
    """Rebuild the contact paragraph with mixed plain-text runs,
    tab characters, and hyperlinks."""
    pe = p._element

    # Capture a template rPr from the first run before nuking children,
    # so the new runs inherit the original font/size formatting.
    first_run = pe.find(qn("w:r"))
    template_rPr = _copy_text_rpr(first_run)

    # Remove all existing <w:r> and <w:hyperlink> children.
    for child in list(pe):
        tag = child.tag.split("}", 1)[-1]
        if tag in ("r", "hyperlink"):
            pe.remove(child)

    # Build new content in order.
    for kind, value in CONTACT_LINE:
        if kind == "text":
            pe.append(_make_text_run(value, template_rPr))
        elif kind == "tab":
            pe.append(_make_tab_run(template_rPr))
        elif kind == "link":
            text, url = value
            pe.append(_make_hyperlink(text, url, template_rPr, header_part))
        else:
            raise ValueError(f"Unknown CONTACT_LINE kind: {kind}")


def trim_dropped_paragraphs(doc) -> int:
    """Remove the paragraphs we marked for deletion. Returns the count
    removed. Idempotent."""
    to_remove = []
    for p in doc.paragraphs:
        stripped = p.text.strip()
        for prefix in PARAGRAPHS_TO_DROP:
            if stripped.startswith(prefix):
                to_remove.append(p)
                break
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return len(to_remove)


def insert_writing_line(doc) -> bool:
    """Append a one-line Writing reference at the end of the document.

    The line lives at the same indent as the role-template paragraphs
    (e.g. 'trader (Python ...)') so it sits as an addendum to PERSONAL
    PROJECTS rather than as a bullet under any specific project.
    Idempotent: returns False if a paragraph starting with 'Writing:'
    is already present."""
    body = doc.paragraphs
    if any(p.text.strip().startswith("Writing:") for p in body):
        return False

    role_template = find_para(
        body, lambda p: p.text.strip().startswith("trader (Python"))
    if role_template is None:
        raise RuntimeError(
            "Could not find role template ('trader (Python...') "
            "— PERSONAL PROJECTS step must run first.")

    # Clone the role template element so we inherit its indent and
    # paragraph properties, then strip its runs so we can rebuild.
    new_el = clone_paragraph_element(role_template)
    for child in list(new_el):
        tag = child.tag.split("}", 1)[-1]
        if tag in ("r", "hyperlink"):
            new_el.remove(child)

    # Capture the role template's first-run formatting (font, size).
    first_run = role_template._element.find(qn("w:r"))
    template_rPr = _copy_text_rpr(first_run)

    # Build prefix text + hyperlink + suffix text.
    new_el.append(_make_text_run(WRITING_LINE["prefix"], template_rPr))
    new_el.append(_make_hyperlink(
        WRITING_LINE["link_text"],
        WRITING_LINE["link_url"],
        template_rPr,
        doc.part,
    ))
    new_el.append(_make_text_run(WRITING_LINE["suffix"], template_rPr))

    # Append after the current last paragraph (which is the last
    # Odds Aggregator bullet after the PERSONAL PROJECTS step has run).
    doc.paragraphs[-1]._element.addnext(new_el)
    return True


# ------------- main -------------

# Fixed sentinel date for scrubbed timestamps — avoids leaking the
# real edit-history dates and keeps re-saves deterministic.
_SCRUB_DATE = datetime(2020, 1, 1, tzinfo=timezone.utc)


def scrub_metadata(doc) -> None:
    """Clear Office core-properties that leak identity / history.

    The committed PDF is the only public artifact; the docx stays local
    after v2.12, but we still scrub on every run so any accidental
    re-tracking of the docx doesn't ship a stranger's name or the real
    edit-history dates. Author is set to the user's own name (already on
    the resume); created/modified are pinned to a fixed sentinel.

    Note: this covers core.xml only. A fully privacy-clean *committed*
    docx would also need app.xml (application/company) and any
    custom.xml scrubbed — see the v2.22 spec's reproducibility section.
    Idempotent."""
    cp = doc.core_properties
    cp.author = "Jake Hoffman"
    cp.last_modified_by = "Jake Hoffman"
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.category = ""
    cp.comments = ""
    cp.identifier = ""
    cp.version = ""
    cp.content_status = ""
    cp.created = _SCRUB_DATE
    cp.modified = _SCRUB_DATE
    cp.revision = 1


def main() -> None:
    doc = Document(str(DOCX))

    # --- 1. Contact line in header (with hyperlinks) ---
    header = doc.sections[0].header
    header_paragraphs = header.paragraphs
    contact_p = find_para(header_paragraphs, lambda p: "@gmail.com" in p.text)
    if contact_p is None:
        raise RuntimeError("Could not find contact line in document header.")

    # Check idempotency: a rebuilt contact line has exactly 3 hyperlinks.
    existing_hlinks = contact_p._element.findall(qn("w:hyperlink"))
    expected_text = "".join(
        v if kind == "text" else
        ("\t" if kind == "tab" else v[0])
        for kind, v in CONTACT_LINE
    )
    # Note: <w:tab/> renders as "\t" when reading paragraph.text.
    if len(existing_hlinks) == 3 and contact_p.text == expected_text:
        print("[same] contact line already has 3 hyperlinks and matching text")
    else:
        build_contact_line(contact_p, header.part)
        print("[ok]   contact line rebuilt with hyperlinks "
              "(email, LinkedIn, website)")

    # --- 2. Skills line in Summary of Qualifications ---
    body = doc.paragraphs
    skills_p = find_para(body, lambda p: p.text.startswith("Proficient in "))
    if skills_p is None:
        raise RuntimeError("Could not find skills paragraph.")
    if skills_p.text != NEW_SKILLS_LINE:
        set_paragraph_text(skills_p, NEW_SKILLS_LINE)
        print("[ok]   skills line updated")
    else:
        print("[same] skills line already updated")

    # --- 3. Trim dropped paragraphs (Lynx bullets + BMC entry) ---
    removed = trim_dropped_paragraphs(doc)
    if removed > 0:
        print(f"[ok]   trimmed {removed} paragraph(s)")
    else:
        print("[same] no paragraphs to trim")

    # --- 4. PERSONAL PROJECTS section ---
    body = doc.paragraphs  # re-read after trim
    experience_heading = find_para(body, lambda p: p.text.strip() == "EXPERIENCE")
    if experience_heading is None:
        raise RuntimeError("Could not find EXPERIENCE heading.")

    role_template = find_para(body, lambda p: p.text.strip().startswith("IT Intern"))
    org_template = find_para(body, lambda p: p.text.strip() == "Lynx Equity Limited")
    bullet_template = find_para(body, lambda p: p.text.startswith("Developed a task management system"))
    if not all([role_template, org_template, bullet_template]):
        raise RuntimeError("Could not find template paragraphs.")

    templates = {
        "heading": experience_heading,
        "role":    role_template,
        "org":     org_template,
        "bullet":  bullet_template,
    }

    # Walk-anchor idempotent insert: for each NEW_SECTION entry, find
    # a matching paragraph in the body. If present, advance the anchor;
    # if absent, clone the template, set text, insert after anchor.
    # Matching: heading uses exact text equality (after strip); role
    # and bullet use prefix match on the first 30 chars (enough to
    # disambiguate without being fragile to em-dash/punctuation drift).
    def _matches(p, kind, text):
        if kind == "heading":
            return p.text.strip() == text
        return p.text.strip().startswith(text[:30].strip())

    anchor_el = doc.paragraphs[-1]._element
    inserted = 0
    skipped = 0
    body = doc.paragraphs  # may have been mutated by trim step
    for kind, text in NEW_SECTION:
        existing = find_para(body, lambda p, k=kind, t=text: _matches(p, k, t))
        if existing is not None:
            anchor_el = existing._element
            skipped += 1
            continue
        template = templates[kind]
        new_el = clone_paragraph_element(template)
        anchor_el.addnext(new_el)
        anchor_el = new_el
        wrapper = Paragraph(new_el, None)
        set_paragraph_text(wrapper, text)
        inserted += 1
        body = doc.paragraphs  # re-read after mutation

    if inserted == 0:
        print(f"[same] PERSONAL PROJECTS section already present ({skipped} entries)")
    else:
        print(f"[ok]   PERSONAL PROJECTS: inserted {inserted} entry(ies), {skipped} already present")

    # --- 5. Writing-section reference appended to PERSONAL PROJECTS ---
    if insert_writing_line(doc):
        print("[ok]   inserted Writing line at end of PERSONAL PROJECTS")
    else:
        print("[same] Writing line already present")

    # --- 4. Scrub Office metadata ---
    scrub_metadata(doc)
    print("[ok]   scrubbed Office metadata (author, comments, etc.)")

    # --- save ---
    doc.save(str(DOCX))
    print(f"wrote {DOCX}")


if __name__ == "__main__":
    main()

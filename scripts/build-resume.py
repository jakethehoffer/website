"""build-resume.py — build resume-source.docx FROM SCRATCH.

Source of truth is tracked text:
  - resume-static.yml : contact header, summary, education, experience
  - projects.yml      : PERSONAL PROJECTS (resume_priority + cap) + the
                        Writing line is appended after them

This replaces the old refresh-resume.py model (which mutated a
gitignored docx via trims + clone-template inserts). The docx is now a
derived artifact: lose it and `python scripts/build-resume.py`
regenerates it from the two tracked YAML files. LibreOffice then
converts the docx to resume.pdf (see README).

Design replicated from the original docx (see scripts/inspect-docx.py):
Garamond; ~0.3in margins; 20pt centered bold name; tab-distributed
contact line with plain-looking hyperlinks; centered bold section
headers; bold role lines with right-aligned dates; italic org lines;
'●' bullets with a hanging indent; 11.5pt body.

Usage: python scripts/build-resume.py
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
STATIC_YML = ROOT / "resume-static.yml"
PROJECTS_YML = ROOT / "projects.yml"
DOCX = ROOT / "resume-source.docx"

# ---- formatting constants (from inspect-docx.py inventory) ----
FONT = "Garamond"
SZ_NAME = 20
SZ_HEADER = 12
SZ_ROLE = 12
SZ_ORG = 12
SZ_BODY = 11.5
MARGIN_LR = 0.3
MARGIN_T = 0.17
MARGIN_B = 0.2
USABLE_WIDTH = 8.5 - 2 * MARGIN_LR          # 7.9in — right tab stop for dates
BULLET_INDENT = 0.27                         # hanging indent for '●' bullets

# Resume shows at most this many projects (top by resume_priority).
RESUME_MAX_PROJECTS = 4

WRITING = {
    "prefix": "Writing: ",
    "link_text": "jakethehoffer.github.io/website/#writing",
    "link_url": "https://jakethehoffer.github.io/website/#writing",
    "suffix": ", short essays on engineering tradeoffs.",
}
_SCRUB_DATE = datetime(2020, 1, 1, tzinfo=timezone.utc)


# ---------------- low-level helpers ----------------

def _style_run(run, *, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def _add_text(paragraph, text, *, size, bold=False, italic=False):
    return _style_run(paragraph.add_run(text), size=size, bold=bold, italic=italic)


def _add_hyperlink(paragraph, text, url, *, size, bold=False):
    """Append a clickable <w:hyperlink> that LOOKS like plain text
    (Garamond, same size, no blue/underline) — matching the original
    resume's contact/writing links."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    hl.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


def _new_paragraph(doc, *, align=None, space_before=0.0, space_after=0.0,
                   line_spacing=1.0, left_indent=None, hanging=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if hanging is not None:
        pf.first_line_indent = Inches(-hanging)
    return p


# ---------------- block renderers ----------------

def render_heading(doc, text):
    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=7.0, space_after=2.0)
    _add_text(p, text, size=SZ_HEADER, bold=True)


def render_role(doc, text, date=None):
    p = _new_paragraph(doc, space_before=3.0, space_after=0.0)
    if date:
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(USABLE_WIDTH), WD_TAB_ALIGNMENT.RIGHT)
        _add_text(p, text, size=SZ_ROLE, bold=True)
        _add_text(p, "\t", size=SZ_ROLE, bold=True)
        _add_text(p, date, size=SZ_ROLE, bold=True)
    else:
        _add_text(p, text, size=SZ_ROLE, bold=True)


def render_org(doc, text):
    p = _new_paragraph(doc, space_after=0.0)
    _add_text(p, text, size=SZ_ORG, italic=True)


def render_bullet(doc, text):
    p = _new_paragraph(doc, space_after=0.0,
                       left_indent=BULLET_INDENT, hanging=BULLET_INDENT)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(BULLET_INDENT))
    _add_text(p, "●", size=SZ_BODY)          # ● filled circle
    _add_text(p, "\t", size=SZ_BODY)
    _add_text(p, text, size=SZ_BODY)


def render_writing_line(doc):
    """Writing pointer with an embedded clickable link. Bold, matching
    the original (it was cloned from a role line)."""
    p = _new_paragraph(doc, space_before=3.0, space_after=0.0)
    _add_text(p, WRITING["prefix"], size=SZ_ROLE, bold=True)
    _add_hyperlink(p, WRITING["link_text"], WRITING["link_url"],
                   size=SZ_ROLE, bold=True)
    _add_text(p, WRITING["suffix"], size=SZ_ROLE, bold=True)


# ---------------- section assembly ----------------

def render_static_block(doc, block):
    kind = block["type"]
    if kind == "role":
        render_role(doc, block["text"], block.get("date"))
    elif kind == "org":
        render_org(doc, block["text"])
    elif kind == "bullets":
        for item in block["items"]:
            render_bullet(doc, item)
    else:
        raise ValueError(f"Unknown static block type: {kind}")


def select_resume_projects(projects):
    """Top RESUME_MAX_PROJECTS candidates (those with a resume block) by
    resume_priority, returned in projects.yml display order."""
    candidates = [p for p in projects if p.get("resume")]
    ranked = sorted(candidates, key=lambda p: p.get("resume_priority", 0),
                    reverse=True)
    keys = {p["key"] for p in ranked[:RESUME_MAX_PROJECTS]}
    return [p for p in projects if p["key"] in keys]


def render_projects_section(doc, projects):
    render_heading(doc, "PERSONAL PROJECTS")
    for p in select_resume_projects(projects):
        r = p["resume"]
        # project role line: bold, no date
        rp = _new_paragraph(doc, space_before=3.0, space_after=0.0)
        _add_text(rp, r["role"], size=SZ_ROLE, bold=True)
        for bullet in r.get("bullets", []):
            render_bullet(doc, bullet)
    render_writing_line(doc)


def build_header(doc, contact):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    # header has one empty paragraph by default; reuse it for the name
    name_p = header.paragraphs[0]
    name_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    name_p.paragraph_format.line_spacing = 1.0
    _style_run(name_p.add_run(contact["name"]), size=SZ_NAME, bold=True)

    # contact line: items left-to-right separated by tabs; links render
    # as plain text (matching the original).
    cp = header.add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.line_spacing = 1.0
    items = contact["line"]
    for i, item in enumerate(items):
        if i > 0:
            _add_text(cp, "\t", size=SZ_BODY)
        if item["type"] == "link":
            _add_hyperlink(cp, item["text"], item["url"], size=SZ_BODY)
        else:
            _add_text(cp, item["text"], size=SZ_BODY)


def scrub(doc):
    cp = doc.core_properties
    cp.author = "Jake Hoffman"
    cp.last_modified_by = "Jake Hoffman"
    cp.title = cp.subject = cp.keywords = cp.category = ""
    cp.comments = cp.identifier = cp.version = cp.content_status = ""
    cp.created = _SCRUB_DATE
    cp.modified = _SCRUB_DATE
    cp.revision = 1


def main():
    static = yaml.safe_load(STATIC_YML.read_text(encoding="utf-8"))
    projects = yaml.safe_load(PROJECTS_YML.read_text(encoding="utf-8"))

    doc = Document()
    # Document defaults: Garamond, body size.
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BODY)

    section = doc.sections[0]
    section.top_margin = Inches(MARGIN_T)
    section.bottom_margin = Inches(MARGIN_B)
    section.left_margin = Inches(MARGIN_LR)
    section.right_margin = Inches(MARGIN_LR)
    section.header_distance = Inches(0.3)

    # Clear any default empty body paragraphs the template ships with,
    # so our content starts clean.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    build_header(doc, static["contact"])

    for sec in static["sections"]:
        render_heading(doc, sec["heading"])
        for block in sec["blocks"]:
            render_static_block(doc, block)

    render_projects_section(doc, projects)

    scrub(doc)
    doc.save(str(DOCX))
    n_proj = len(select_resume_projects(projects))
    print(f"[ok]   built {DOCX.name} from resume-static.yml + projects.yml "
          f"({n_proj} projects)")


if __name__ == "__main__":
    main()
